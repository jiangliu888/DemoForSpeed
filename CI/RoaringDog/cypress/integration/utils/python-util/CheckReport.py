# -*- coding: UTF-8 -*-
import os
import json
import urlparse
from docx import Document
from CheckConsul import call_funcion
from collections import OrderedDict

class ReportCheck:
    def __init__(self):
        self._dir = os.path.abspath(os.path.join(os.path.abspath(__file__), "../../../../downloads"))
        if not os.path.exists(self._dir):
            os.mkdir(self._dir)
        _file = ''
        for f in os.listdir(self._dir):
            if f:
                _file = f
                break
        if _file:
            self.file = Document(self._dir + '/' + _file)
            self.tables = self.file.tables

    def export_success_check(self):
        flag = 0
        for f in os.listdir(self._dir):
            if f:
                flag = 1
                break
        assert flag == 1

    
    def empty_downloads_dir(self):
        if os.path.exists(self._dir):
            for f in os.listdir(self._dir):
                if f:
                    os.remove(self._dir + '/' + f)
        else:
            os.mkdir(self._dir)

    def customer_name_check(self, argList):
        company_name = argList[0]
        assert self.file.paragraphs[3].text == '客户名称：' + company_name,'{0} {1}'.format(self.file.paragraphs[3].text,'客户名称：' + company_name)

    def service_overview_check(self, argList):
        x = str(argList[0])
        value = x.replace('%u','\\u') #将%uxxxx 替换换 \uxxxx 这才可以进行utf-8解码
        byts = urlparse.unquote(value) #返回的 byte
        #byts = byts.decode('UTF-8') # decode UTF-8 解码只能解开 \uXXXX 的Unicode 标准形式
        data = json.loads(byts,object_pairs_hook=OrderedDict)            
        keys = list(data.keys())
        values = list(data.values())
        table = self.tables[0]
        assert len(data.keys()) == len(table.rows)
        for row in range(len(table.rows)):
            assert len(table.rows[row].cells) == 2
            assert table.cell(row, 0).text.replace('\n','') == keys[row],'{0} {1}'.format(table.cell(row,0).text.replace('\n',''),keys[row])
            assert table.cell(row, 1).text.replace('\n','') == values[row],'{0} {1}'.format(table.cell(row,1).text.replace('\n',''),values[row])
    
    def service_quality_check(self,argList):
        x = str(argList[0])
        value = x.replace('%u','\\u') #将%uxxxx 替换换 \uxxxx 这才可以进行utf-8解码
        byts = urlparse.unquote(value) #返回的 byte
        #byts = byts.decode('UTF-8') # decode UTF-8 解码只能解开 \uXXXX 的Unicode 标准形式
        data = json.loads(byts,object_pairs_hook=OrderedDict)            
        table = self.tables[2]
        #ll = len(table.rows[0].cells)
        keys = list(data.keys())
        assert len(data.keys()) == len(table.rows)
        for row in range(len(data)):
            assert table.cell(row, 0).text.replace('\n', '') == keys[row],'{0} {1}'.format(table.cell(row, 0).text.replace('\n', ''),keys[row])
            for j in range(len(data[keys[row]])):
                assert table.cell(row, j+1).text.replace('\n', '') == data[keys[row]][j]              
        
    def service_item_check(self, argList):
        item = argList[0]
        expect = argList[1]
        table = self.tables[0]
        for row in range(1, len(table.rows)):
            if table.cell(row, 0) == item:
                assert table.cell(row, 1).text == expect
                break
        
    def device_item_check(self, argList):
        device = argList[0]
        item = argList[1]
        expect = argList[2]
        table = self.tables[1]
        columns = len(table.rows[0].cells)
        for i in range(columns):
            if table.cell(0, i).text.replace('\n', '') == item:
                c = i
                break 
        for row in range(1, len(table.rows)):
            if table.cell(row, 0).text.replace('\n', '') == device:
                assert table.cell(row, c).text.replace('\n', '') == expect

    def device_status_check(self, argList):
        x = str(argList[0])
        value = x.replace('%u','\\u') #将%uxxxx 替换换 \uxxxx 这才可以进行utf-8解码
        byts = urlparse.unquote(value) #返回的 byte
        #byts = byts.decode('UTF-8') # decode UTF-8 解码只能解开 \uXXXX 的Unicode 标准形式
        data = json.loads(byts,object_pairs_hook=OrderedDict)            
        table = self.tables[1]
        #ll = len(table.rows[0].cells)
        keys = list(data.keys())
        assert len(data.keys()) == len(table.rows)
        for row in range(len(data)):
            assert table.cell(row, 0).text.replace('\n', '') == keys[row],'{0} {1}'.format(table.cell(row,0).text.replace('\n',''),keys[row])
            for j in range(len(data[keys[row]])):
                assert table.cell(row, j+1).text.replace('\n', '') == data[keys[row]][j]

    def ticket_detail_check(self,argList):
        x=argList[0]
        value = x.replace('%u','\\u')
        byts = urlparse.unquote(value) #返回的 byte
        #byts = byts.decode('UTF-8') # decode UTF-8 解码只能解开 \uXXXX 的Unicode 标准形式
        data = json.loads(byts,object_pairs_hook=OrderedDict)
        table = self.tables[2]
        keys = list(data.keys())
        assert len(keys) == len(table.rows)
        for row in range(len(data)):
            assert table.cell(row, 0).text.replace('\n', '') == keys[row],'{0} {1}'.format(table.cell(row,0).text.replace('\n',''),keys[row])
            if len(data)>2 and row == 1:
                assert table.cell(row, 2).text.replace('\n', '') == data[keys[row]][0]
                assert table.cell(row, 4).text.replace('\n', '') == data[keys[row]][1]
            else:
                for j in range(len(data[keys[row]])):                
                    assert table.cell(row, j+1).text.replace('\n', '') == data[keys[row]][j], '{0}\n{1}'.format(table.cell(row, j+1).text.replace('\n', ''),data[keys[row]][j])
            
         
if __name__ == "__main__":
    
    match = \
        {
            'customer_name_check': ReportCheck().customer_name_check,
            'service_overview_check': ReportCheck().service_overview_check,
            'service_item_check': ReportCheck().service_item_check,
            'device_item_check': ReportCheck().device_item_check,
            'device_status_check': ReportCheck().device_status_check,
            'empty_downloads_dir': ReportCheck().empty_downloads_dir,
            'export_success_check': ReportCheck().export_success_check,
            'service_quality_check': ReportCheck().service_quality_check,
            'ticket_detail_check': ReportCheck().ticket_detail_check
        }
    call_funcion(match)
